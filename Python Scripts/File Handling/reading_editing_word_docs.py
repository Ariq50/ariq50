import docx

d = docx.Document('C:\\Users\\Sophytes\\Documents\\Misc\\demo.docx')

p = d.paragraphs[0].text

p.runs[0].italic = True

p.runs[0].underline = True

p.runs[0].text = 'italic and underline'

p.style = 'Title'

d.save('C:\\Users\\Sophytes\\Documents\\Misc\\demo2.docx)

d2 = docx.Document()

d2.add_paragraph('Boopboopbeepboop!')

d2.save('C:\\Users\\Sophytes\\Documents\\Misc\\demo3.docx')

